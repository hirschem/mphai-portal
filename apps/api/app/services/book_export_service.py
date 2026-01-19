from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path


class BookExportService:
    """Service for exporting book chapters to Word documents"""
    
    def export_chapter(self, chapter_name: str, transcribed_text: str, output_path: Path):
        """Export chapter as editable Word document"""
        
        doc = Document()
        
        # Add chapter title
        title = doc.add_heading(chapter_name, level=1)
        title.style.font.size = Pt(18)
        
        doc.add_paragraph()  # Spacing
        
        # Split text into pages (if marked with [Page X])
        pages = transcribed_text.split('[Page ')
        
        for page_section in pages:
            if not page_section.strip():
                continue
            
            # Remove page marker and get content
            if ']' in page_section:
                page_content = page_section.split(']', 1)[1].strip()
            else:
                page_content = page_section.strip()
            
            # Add content paragraphs
            paragraphs = page_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.style.font.size = Pt(12)
                    para.style.font.name = 'Times New Roman'
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Save document
        doc.save(str(output_path))
